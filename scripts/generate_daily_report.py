#!/usr/bin/env python3
"""生成每日交易报告（Markdown + 可选 DOCX）。"""

from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import List

from _config import PROJECT_ROOT


def _load_trade_files(log_dir: Path, day: str) -> List[dict]:
    rows = []
    for fp in sorted(log_dir.glob(f"trade_{day.replace('-', '')}_*.json")):
        try:
            data = json.loads(fp.read_text(encoding="utf-8"))
            data["_file"] = str(fp)
            rows.append(data)
        except Exception:
            continue
    return rows


def _sum_receipts(trades: List[dict]) -> dict:
    submitted = 0
    failed = 0
    skipped = 0
    buy_amt = 0.0
    sell_amt = 0.0
    for t in trades:
        recs = (t.get("execution_result", {}) or {}).get("receipts", []) or []
        for r in recs:
            status = str(r.get("status", ""))
            side = str(r.get("side", "")).lower()
            amt = float(r.get("order_amount_cny", 0) or 0)
            if status in {"submitted", "filled", "simulated"}:
                submitted += 1
            elif status == "failed":
                failed += 1
            else:
                skipped += 1
            if side == "buy":
                buy_amt += amt
            elif side == "sell":
                sell_amt += amt
    return {
        "submitted": submitted,
        "failed": failed,
        "skipped": skipped,
        "buy_amt": round(buy_amt, 2),
        "sell_amt": round(sell_amt, 2),
    }


def _to_markdown(day: str, trades: List[dict], s: dict) -> str:
    lines = [
        f"# QuantLLM 日报 {day}",
        "",
        "## 摘要",
        f"- 交易批次: {len(trades)}",
        f"- 成功/模拟回执: {s['submitted']}",
        f"- 失败回执: {s['failed']}",
        f"- 跳过回执: {s['skipped']}",
        f"- 买入金额(元): {s['buy_amt']}",
        f"- 卖出金额(元): {s['sell_amt']}",
        "",
        "## 交易文件",
    ]
    for t in trades[-20:]:
        lines.append(f"- `{Path(t.get('_file', '')).name}` mode={t.get('mode', '')} regime={t.get('market_regime', '')}")
    lines.append("")
    lines.append("## 备注")
    lines.append("- 本报告由脚本自动生成，仅供复盘。")
    return "\n".join(lines) + "\n"


def _save_docx(md_text: str, fp: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False
    doc = Document()
    for ln in md_text.splitlines():
        if ln.startswith("# "):
            doc.add_heading(ln[2:], level=1)
        elif ln.startswith("## "):
            doc.add_heading(ln[3:], level=2)
        elif ln.startswith("- "):
            doc.add_paragraph(ln[2:], style="List Bullet")
        else:
            doc.add_paragraph(ln)
    doc.save(str(fp))
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate daily trade report")
    parser.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"), help="报告日期 YYYY-MM-DD")
    parser.add_argument("--log-dir", default="output/trade_logs", help="交易日志目录")
    parser.add_argument("--out-dir", default="docs/daily_reports", help="报告输出目录")
    parser.add_argument("--copy-to", default="", help="可选：复制 Markdown 到指定目录（如 ~/Desktop/photo）")
    args = parser.parse_args()

    log_dir = Path(PROJECT_ROOT) / args.log_dir
    out_dir = Path(PROJECT_ROOT) / args.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    trades = _load_trade_files(log_dir, args.date)
    stat = _sum_receipts(trades)
    md = _to_markdown(args.date, trades, stat)

    base = f"daily_report_{args.date.replace('-', '')}"
    md_fp = out_dir / f"{base}.md"
    docx_fp = out_dir / f"{base}.docx"
    md_fp.write_text(md, encoding="utf-8")
    has_docx = _save_docx(md, docx_fp)

    if args.copy_to.strip():
        target = Path(args.copy_to).expanduser()
        target.mkdir(parents=True, exist_ok=True)
        (target / md_fp.name).write_text(md, encoding="utf-8")

    print(f"[REPORT] trades={len(trades)} markdown={md_fp}")
    print(f"[REPORT] docx={'ok:'+str(docx_fp) if has_docx else 'skipped:python-docx not installed'}")


if __name__ == "__main__":
    main()
